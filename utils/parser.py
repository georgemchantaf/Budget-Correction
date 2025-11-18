import docx
import openpyxl
from pypdf import PdfReader
import re
import pandas as pd
from typing import Dict, List, Any

class DocumentParser:
    """Parse Word, Excel, and PDF documents to extract budget data"""
    
    def __init__(self):
        self.supported_formats = ['.docx', '.xlsx', '.pdf']
    
    def parse(self, uploaded_file) -> Dict[str, Any]:
        """Main parsing function - routes to appropriate parser"""
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        if file_extension == 'docx':
            return self.parse_word(uploaded_file)
        elif file_extension == 'xlsx':
            return self.parse_excel(uploaded_file)
        elif file_extension == 'pdf':
            return self.parse_pdf(uploaded_file)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def parse_word(self, uploaded_file) -> Dict[str, Any]:
        """Parse Word document"""
        doc = docx.Document(uploaded_file)
        
        # Extract student info
        full_text = '\n'.join([para.text for para in doc.paragraphs])
        student_name = self._extract_student_name(full_text)
        department = self._extract_department(full_text)
        
        # Extract tables
        tables_data = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables_data.append(table_data)
        
        # Parse budget data from tables
        parsed_data = self._parse_budget_tables(tables_data)
        
        return {
            'student_name': student_name,
            'department': department,
            **parsed_data
        }
    
    def parse_excel(self, uploaded_file) -> Dict[str, Any]:
        """Parse Excel document"""
        wb = openpyxl.load_workbook(uploaded_file)
        ws = wb.active
        
        # Extract all data
        all_data = []
        for row in ws.iter_rows(values_only=True):
            all_data.append(list(row))
        
        # Extract student info from first few rows
        full_text = ' '.join([str(cell) for row in all_data[:10] for cell in row if cell])
        student_name = self._extract_student_name(full_text)
        department = self._extract_department(full_text)
        
        # Parse budget tables
        parsed_data = self._parse_budget_tables([all_data])
        
        return {
            'student_name': student_name,
            'department': department,
            **parsed_data
        }
    
    def parse_pdf(self, uploaded_file) -> Dict[str, Any]:
        """Parse PDF document"""
        pdf_reader = PdfReader(uploaded_file)
        
        # Extract text from all pages
        full_text = ''
        for page in pdf_reader.pages:
            full_text += page.extract_text()
        
        student_name = self._extract_student_name(full_text)
        department = self._extract_department(full_text)
        
        # For PDFs, we'll rely more on AI extraction since table parsing is harder
        # Return raw text for AI to process
        return {
            'student_name': student_name,
            'department': department,
            'raw_text': full_text,
            'needs_ai_extraction': True
        }
    
    def _extract_student_name(self, text: str) -> str:
        """Extract student name from text"""
        # Common patterns
        patterns = [
            r'(?:Student|Name|By|Author):\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)',
            r'^([A-Z][a-z]+\s+[A-Z][a-z]+)',  # First line with capitalized name
            r'([A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\s*\n',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.MULTILINE)
            if match:
                return match.group(1).strip()
        
        return "Unknown Student"
    
    def _extract_department(self, text: str) -> str:
        """Extract department/unit from text"""
        # Look for department keywords
        dept_keywords = [
            'Emergency Department', 'ED', 'ER', 'Emergency Room',
            'Neonatal Intensive Care', 'NICU', 'ICU',
            'Pediatric', 'Surgical', 'Medical',
            'Nursing', 'HSON', 'Hariri School'
        ]
        
        for keyword in dept_keywords:
            if keyword.lower() in text.lower():
                # Try to extract full context
                pattern = rf'([^.]*{re.escape(keyword)}[^.]*)'
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    return match.group(1).strip()[:100]  # Limit length
        
        return "Unknown Department"
    
    def _parse_budget_tables(self, tables_data: List[List[List[str]]]) -> Dict[str, Any]:
        """Parse budget tables to extract structured data"""
        
        result = {
            'fixed_expenses': [],
            'variable_expenses': [],
            'total_expenses': {},
            'patient_days_initial': None
        }
        
        for table in tables_data:
            if not table or len(table) < 2:
                continue
            
            # Identify table type by headers
            headers = [str(cell).lower().strip() for cell in table[0]]
            
            # Check if it's Fixed Expenses table
            # Look for: "fixed" keyword OR ("monthly" AND "2024" columns)
            is_fixed = (any('fixed' in h for h in headers) or
                       (any('monthly' in h for h in headers) and
                        any('2024' in h for h in headers)))

            # Check if it's Variable Expenses table
            # Look for: "variable" keyword OR ("pt day" OR "patient day" columns)
            is_variable = (any('variable' in h for h in headers) or
                          any('pt day' in h or 'patient day' in h for h in headers))

            # Check if it's Total Expenses table
            # Look for: ("total" AND column headers like "yearly total" or "inflation")
            is_total = (any('total' in h for h in headers) and
                       any('yearly' in h or 'inflation' in h for h in headers) and
                       len(table) <= 3)  # Total table is usually 1-2 rows

            if is_fixed:
                result['fixed_expenses'] = self._parse_fixed_table(table)
            elif is_variable:
                result['variable_expenses'] = self._parse_variable_table(table)
            elif is_total:
                result['total_expenses'] = self._parse_total_table(table)
            
            # Check for initial patient days table
            elif any('patient days' in str(cell).lower() for row in table for cell in row):
                result['patient_days_initial'] = self._extract_patient_days(table)
        
        return result
    
    def _parse_fixed_table(self, table: List[List[str]]) -> List[Dict[str, Any]]:
        """Parse fixed expenses table with flexible column detection"""
        fixed_items = []

        if not table or len(table) < 2:
            return fixed_items

        # Get headers and find column indices
        headers = [str(cell).lower().strip() for cell in table[0]]

        # Find column indices by matching keywords
        # Process in priority order: check 2025 before 2024 to avoid conflicts
        col_map = {}
        for i, h in enumerate(headers):
            if any(kw in h for kw in ['description', 'expense', 'item']):
                col_map['description'] = i
            elif '5' in h and 'month' in h:
                col_map['5_month'] = i
            elif 'monthly' in h and '2024' not in h and '2025' not in h and 'year' not in h:
                col_map['monthly'] = i
            elif '2025' in h or 'estimate' in h:
                # Check 2025/estimate BEFORE 2024 to avoid wrong mapping
                col_map['2025_estimate'] = i
            elif '2024' in h:
                # Only match if 2024 is explicitly in header
                col_map['2024_year'] = i
            elif 'inflation' in h and ('rate' in h or '%' in h):
                col_map['inflation_rate'] = i
            elif 'inflation' in h and ('amount' in h or '$' in h or 'dollar' in h):
                col_map['inflation_amount'] = i

        # Parse data rows
        for row in table[1:]:
            if not row or len(row) < 3:
                continue

            try:
                item = {
                    'description': self._clean_text(row[col_map.get('description', 0)]),
                    '5_month_consumption': self._parse_number(row[col_map['5_month']]) if '5_month' in col_map else None,
                    'monthly_consumption': self._parse_number(row[col_map['monthly']]) if 'monthly' in col_map else None,
                    '2024_year_consumption': self._parse_number(row[col_map['2024_year']]) if '2024_year' in col_map else None,
                    'inflation_rate': self._parse_number(row[col_map['inflation_rate']]) if 'inflation_rate' in col_map else None,
                    'inflation_amount': self._parse_number(row[col_map['inflation_amount']]) if 'inflation_amount' in col_map else None,
                    'estimated_2025_consumption': self._parse_number(row[col_map['2025_estimate']]) if '2025_estimate' in col_map else None
                }

                # Skip subtotal and total rows (use word boundaries)
                desc_lower = item['description'].lower()
                # Check for whole words only, not substrings
                import re
                skip_pattern = r'\b(total|subtotal|sum)\b'
                if item['description'] and not re.search(skip_pattern, desc_lower):
                    fixed_items.append(item)
            except (ValueError, IndexError, KeyError):
                continue

        return fixed_items
    
    def _parse_variable_table(self, table: List[List[str]]) -> List[Dict[str, Any]]:
        """Parse variable expenses table with flexible column detection"""
        variable_items = []

        if not table or len(table) < 2:
            return variable_items

        # Get headers and find column indices
        headers = [str(cell).lower().strip() for cell in table[0]]

        # Find column indices by matching keywords
        # Be specific: check more specific conditions first
        col_map = {}
        for i, h in enumerate(headers):
            # Description column
            if any(kw in h for kw in ['description', 'item']) and not any(x in h for x in ['consumption', 'amount', 'day']):
                col_map['description'] = i
            elif 'expense' in h and len(h) < 15 and not any(x in h for x in ['consumption', 'cons.', 'amount', 'day']):
                # "Expense" by itself (short header)
                col_map['description'] = i
            # 5-month consumption: "5-month consumption" OR "5-month Cons."
            elif ('5-month' in h or '5m' in h) and ('consumption' in h or 'cons' in h):
                col_map['5_month_cons'] = i
            # 5-month patient days: "5-month patient days" OR "5m Pt Days"
            elif ('5-month' in h or '5m' in h) and ('pt' in h or 'patient') and 'day' in h:
                col_map['5_month_days'] = i
            # Consumption per patient day: "consumption per patient day" OR "Cons per Pt Day"
            elif ('consumption' in h or 'cons' in h) and 'per' in h and ('pt' in h or 'patient') and 'day' in h:
                col_map['cons_per_day'] = i
            # Estimated yearly pt days: "Estimated 2025 yearly pt. days" OR "2025 Pt Days"
            elif ('2025' in h or 'estimated' in h or 'yearly' in h) and ('pt' in h or 'patient') and 'day' in h and 'amount' not in h:
                col_map['yearly_days'] = i
            # Amount per yearly pt days: "Amount per yearly pt. days" OR "Yearly Amount"
            elif ('yearly' in h or 'per' in h) and 'amount' in h:
                col_map['yearly_amount'] = i
            # Inflation rate
            elif 'inflation' in h and ('rate' in h or '%' in h):
                col_map['inflation_rate'] = i
            # Inflation amount
            elif 'inflation' in h and ('amount' in h or '$' in h):
                col_map['inflation_amount'] = i
            # Total amount: "Total amount" OR "Total 2025"
            elif 'total' in h and ('amount' in h or '2025' in h):
                col_map['total'] = i

        # Parse data rows
        for row in table[1:]:
            if not row or len(row) < 3:
                continue

            try:
                item = {
                    'description': self._clean_text(row[col_map.get('description', 0)]),
                    '5_month_consumption': self._parse_number(row[col_map['5_month_cons']]) if '5_month_cons' in col_map else None,
                    '5_month_patient_days': self._parse_number(row[col_map['5_month_days']]) if '5_month_days' in col_map else None,
                    'consumption_per_patient_day': self._parse_number(row[col_map['cons_per_day']]) if 'cons_per_day' in col_map else None,
                    'estimated_2025_yearly_pt_days': self._parse_number(row[col_map['yearly_days']]) if 'yearly_days' in col_map else None,
                    'amount_per_yearly_pt_days': self._parse_number(row[col_map['yearly_amount']]) if 'yearly_amount' in col_map else None,
                    'inflation_rate': self._parse_number(row[col_map['inflation_rate']]) if 'inflation_rate' in col_map else None,
                    'inflation_amount': self._parse_number(row[col_map['inflation_amount']]) if 'inflation_amount' in col_map else None,
                    'total_amount': self._parse_number(row[col_map['total']]) if 'total' in col_map else None
                }

                # Skip subtotal and total rows (use word boundaries)
                desc_lower = item['description'].lower()
                # Check for whole words only, not substrings
                import re
                skip_pattern = r'\b(total|subtotal|sum)\b'
                if item['description'] and not re.search(skip_pattern, desc_lower):
                    variable_items.append(item)
            except (ValueError, IndexError, KeyError):
                continue

        return variable_items
    
    def _parse_total_table(self, table: List[List[str]]) -> Dict[str, Any]:
        """Parse total expenses table with flexible column detection"""
        if not table or len(table) < 2:
            return {}

        # Get headers and find column indices
        headers = [str(cell).lower().strip() for cell in table[0]]

        # Find column indices by matching keywords
        col_map = {}
        for i, h in enumerate(headers):
            if '5-month' in h or ('5' in h and 'month' in h and i < 3):
                # "5-month" column (usually first data column)
                col_map['5_month'] = i
            elif 'yearly' in h and ('consumption' in h or 'total' in h):
                # "Yearly consumption" or "Yearly Total"
                col_map['yearly'] = i
            elif 'inflation' in h and ('rate' in h or '%' in h):
                col_map['inflation_rate'] = i
            elif 'inflation' in h and ('amount' in h or '$' in h):
                col_map['inflation_amount'] = i
            elif ('total' in h and ('2025' in h or 'amount' in h)) or (i == len(headers) - 1 and 'total' in h):
                # "Total amount" or "Total 2025" (usually last column)
                col_map['total'] = i

        # Parse the data row (usually just one row)
        for row in table[1:]:
            if not row or len(row) < 3:
                continue

            try:
                return {
                    '5_month_consumption': self._parse_number(row[col_map['5_month']]) if '5_month' in col_map else None,
                    'yearly_consumption': self._parse_number(row[col_map['yearly']]) if 'yearly' in col_map else None,
                    'inflation_rate': self._parse_number(row[col_map['inflation_rate']]) if 'inflation_rate' in col_map else None,
                    'inflation_amount': self._parse_number(row[col_map['inflation_amount']]) if 'inflation_amount' in col_map else None,
                    'total_amount': self._parse_number(row[col_map['total']]) if 'total' in col_map else None
                }
            except (ValueError, IndexError, KeyError):
                continue

        return {}
    
    def _extract_patient_days(self, table: List[List[str]]) -> int:
        """Extract patient days from initial table"""
        for row in table:
            for i, cell in enumerate(row):
                if 'patient days' in str(cell).lower():
                    # Next cell should have the value
                    if i + 1 < len(row):
                        return self._parse_number(row[i + 1])
        return None
    
    def _clean_text(self, text: str) -> str:
        """Clean text from table cells"""
        if not text:
            return ""
        # Remove extra whitespace, newlines
        cleaned = ' '.join(str(text).split())
        # Remove special characters but keep alphanumeric and basic punctuation
        cleaned = re.sub(r'[^\w\s\-/(),.]', '', cleaned)
        return cleaned.strip()
    
    def _parse_number(self, value: Any) -> float:
        """Parse numeric value from string, including formulas like '500/5= 100$'"""
        if value is None or value == '' or str(value).lower() == 'none':
            return None

        # Convert to string and clean
        value_str = str(value).strip()

        # Check if empty after stripping
        if not value_str or value_str == '':
            return None

        # Check if it contains a formula (e.g., "500/5= 100$" or "100 x 12= 1,200$")
        # Extract the result after the equals sign
        if '=' in value_str:
            # Split by = and take the part after it
            parts = value_str.split('=')
            if len(parts) >= 2:
                # Take the last part (the result)
                value_str = parts[-1].strip()

        # Remove currency symbols, commas, percentage signs, spaces
        value_str = re.sub(r'[$,\s%]', '', value_str)

        # Check again after cleaning
        if not value_str or value_str == '':
            return None

        # Try to convert to float
        try:
            return float(value_str)
        except (ValueError, TypeError):
            return None